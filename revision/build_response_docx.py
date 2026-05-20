from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "MBE_response_to_reviewers_with_answers.md"
OUTPUT = ROOT / "MBE_response_to_reviewers_with_answers.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_runs_from_inline_markdown(paragraph, text):
    # Handles the limited inline Markdown used by the response source:
    # bold labels and backtick code spans.
    tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(token)


def add_markdown_paragraph(document, line):
    if line.startswith("**Comment"):
        p = document.add_paragraph(style="Comment")
    elif line.startswith("**Response"):
        p = document.add_paragraph(style="Response")
    else:
        p = document.add_paragraph(style="Normal")
    add_runs_from_inline_markdown(p, line)
    return p


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    subtitle.paragraph_format.space_after = Pt(14)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    if "Comment" not in styles:
        styles.add_style("Comment", 1)
    comment = styles["Comment"]
    comment.base_style = normal
    comment.font.italic = True
    comment.font.color.rgb = RGBColor(0x43, 0x43, 0x43)
    comment.paragraph_format.left_indent = Inches(0.18)
    comment.paragraph_format.space_before = Pt(4)
    comment.paragraph_format.space_after = Pt(4)
    comment.paragraph_format.line_spacing = 1.1

    if "Response" not in styles:
        styles.add_style("Response", 1)
    response = styles["Response"]
    response.base_style = normal
    response.paragraph_format.left_indent = Inches(0.18)
    response.paragraph_format.space_before = Pt(2)
    response.paragraph_format.space_after = Pt(8)
    response.paragraph_format.line_spacing = 1.1


def add_footer(document):
    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    run = footer.add_run("Phylo-Movies revision response")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build():
    markdown = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    add_footer(doc)

    first_body_paragraph = True
    for line in markdown:
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(line[2:], style="Title")
            set_paragraph_spacing(p, after=4)
        elif line.startswith("Manuscript:") or line.startswith("Date:"):
            p = doc.add_paragraph(style="Subtitle")
            add_runs_from_inline_markdown(p, line)
        elif line.startswith("## "):
            if not first_body_paragraph:
                doc.add_paragraph()
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        else:
            add_markdown_paragraph(doc, line)
            first_body_paragraph = False

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
